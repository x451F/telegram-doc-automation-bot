"""Generic DOCX generation utilities for workflow payload output."""

from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
from typing import Mapping, TYPE_CHECKING

from app.services.payloads import DocumentIntakePayload
from app.services.template_mapping import build_template_mapping
from app.services.text_utils import safe_filename

if TYPE_CHECKING:
    from docx.document import Document as DocxDocument
    from docx.table import Table
    from docx.text.paragraph import Paragraph


class DocumentGenerationError(RuntimeError):
    """Raised when template rendering fails."""


@dataclass(slots=True, frozen=True)
class GeneratedDocumentPaths:
    """Filesystem paths to generated DOCX output."""

    contract_path: Path
    completion_certificate_path: Path


@dataclass(slots=True, frozen=True)
class TemplateFiles:
    """Expected input template files."""

    contract_template_path: Path
    completion_certificate_template_path: Path


def resolve_template_files(templates_dir: Path) -> TemplateFiles:
    """Resolve expected template file paths from configured directory."""
    return TemplateFiles(
        contract_template_path=templates_dir / "contract_template.docx",
        completion_certificate_template_path=templates_dir / "completion_certificate_template.docx",
    )


def build_output_filename(prefix: str, number: str, date_value: str) -> str:
    """Build deterministic output filename based on stable payload fields."""
    stem = safe_filename(f"{prefix}-{number}-{date_value}")
    return f"{stem}.docx"


def _replace_text(text: str, placeholders: Mapping[str, str]) -> str:
    rendered = text
    for key in sorted(placeholders.keys(), key=len, reverse=True):
        rendered = rendered.replace(key, placeholders[key])
    return rendered


def _replace_placeholders_in_paragraph(
    paragraph: Paragraph,
    placeholders: Mapping[str, str],
) -> None:
    if not paragraph.text:
        return
    paragraph.text = _replace_text(paragraph.text, placeholders)


def _replace_placeholders_in_table(
    table: Table,
    placeholders: Mapping[str, str],
) -> None:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                _replace_placeholders_in_paragraph(paragraph, placeholders)


def _render_document(document: DocxDocument, placeholders: Mapping[str, str]) -> None:
    for paragraph in document.paragraphs:
        _replace_placeholders_in_paragraph(paragraph, placeholders)
    for table in document.tables:
        _replace_placeholders_in_table(table, placeholders)


def render_docx_template(
    template_path: Path,
    output_path: Path,
    placeholders: Mapping[str, str],
) -> Path:
    """Render single DOCX template to output path using placeholder mapping."""
    if not template_path.exists():
        raise FileNotFoundError(f"Template file does not exist: {template_path}")

    try:
        from docx import Document
    except ModuleNotFoundError as exc:
        raise DocumentGenerationError(
            "python-docx is required to render templates. Install project dependencies first."
        ) from exc

    try:
        document = Document(str(template_path))
        _render_document(document, placeholders)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        document.save(str(output_path))
    except Exception as exc:
        raise DocumentGenerationError(f"Failed to render template: {template_path}") from exc

    return output_path


def generate_documents(
    payload: DocumentIntakePayload,
    templates_dir: Path,
    output_dir: Path,
    city: str | None = None,
) -> GeneratedDocumentPaths:
    """Generate agreement and completion certificate DOCX files."""
    template_files = resolve_template_files(templates_dir)
    mapping = build_template_mapping(payload, city=city)

    contract_filename = build_output_filename(
        prefix="service-agreement",
        number=payload.contract_number,
        date_value=payload.contract_date,
    )
    certificate_filename = build_output_filename(
        prefix="completion-certificate",
        number=payload.certificate_number,
        date_value=payload.certificate_date,
    )

    contract_output_path = output_dir / contract_filename
    certificate_output_path = output_dir / certificate_filename

    render_docx_template(
        template_path=template_files.contract_template_path,
        output_path=contract_output_path,
        placeholders=mapping.contract,
    )
    render_docx_template(
        template_path=template_files.completion_certificate_template_path,
        output_path=certificate_output_path,
        placeholders=mapping.completion_certificate,
    )

    return GeneratedDocumentPaths(
        contract_path=contract_output_path,
        completion_certificate_path=certificate_output_path,
    )

